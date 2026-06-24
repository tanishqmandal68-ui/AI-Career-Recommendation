"""Document extraction utilities for PDF, DOCX, text, and image resume inputs."""

from __future__ import annotations

import base64
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
import re
from typing import Any, BinaryIO, Dict, Iterable, List, Optional

from docx import Document
import pdfplumber

from backend.core.api_manager import APIManager
from backend.core.logger import get_logger

LOGGER = get_logger(__name__)
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
DOC_EXTENSIONS = {".doc", ".docx"}
SUPPORTED_EXTENSIONS = {".pdf", *DOC_EXTENSIONS, ".txt", *IMAGE_EXTENSIONS}


@dataclass
class ParsedDocument:
    """Structured extraction result for uploaded career documents."""

    text: str
    extension: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    warnings: List[str] = field(default_factory=list)


def extract_text(
    file_obj: BinaryIO,
    filename: str,
    api_manager: Optional[APIManager] = None,
) -> ParsedDocument:
    """Extract text from a supported document or image file.

    Args:
        file_obj: Binary file object from disk or Streamlit upload.
        filename: Original filename used to detect the extension.
        api_manager: Optional API manager used for Kimi Vision image extraction.

    Returns:
        Parsed document with text, metadata, and warnings.

    Raises:
        ValueError: If the extension is unsupported or extraction cannot produce text.
    """
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {extension or 'unknown'}")
    payload = file_obj.read()
    if not payload:
        raise ValueError("Uploaded file is empty.")
    if extension == ".pdf":
        return extract_pdf_text(payload)
    if extension == ".docx":
        return extract_docx_text(payload)
    if extension == ".doc":
        return extract_doc_text(payload)
    if extension == ".txt":
        return ParsedDocument(text=normalize_text(payload.decode("utf-8", errors="replace")), extension=extension)
    return extract_image_text(payload, extension, api_manager)


def extract_pdf_text(payload: bytes) -> ParsedDocument:
    """Extract text and layout warnings from a PDF payload."""
    text_parts: List[str] = []
    warnings: List[str] = []
    metadata: Dict[str, Any] = {}
    with pdfplumber.open(BytesIO(payload)) as pdf:
        metadata = {str(key): str(value) for key, value in (pdf.metadata or {}).items() if value is not None}
        metadata["pages"] = len(pdf.pages)
        for page_number, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text(layout=True) or page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)
            words = page.extract_words() or []
            if len(words) < 8 and page.images:
                warnings.append(f"Page {page_number} appears image-heavy and may be scanned.")
            if has_multicolumn_layout(words):
                warnings.append(f"Page {page_number} may use a multi-column layout.")
            if len(page.chars) > 0 and not page_text.strip():
                warnings.append(f"Page {page_number} contains text objects that were not cleanly extracted.")
            if page.images and len(page.images) > 2:
                warnings.append(f"Page {page_number} contains multiple graphics.")
    text = normalize_text("\n".join(text_parts))
    if not text:
        raise ValueError("No extractable text found in the PDF.")
    hidden_warning = detect_hidden_text(text)
    if hidden_warning:
        warnings.append(hidden_warning)
    return ParsedDocument(text=text, extension=".pdf", metadata=metadata, warnings=dedupe(warnings))


def extract_docx_text(payload: bytes) -> ParsedDocument:
    """Extract paragraph and table text from a DOCX payload."""
    document = Document(BytesIO(payload))
    parts: List[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    text = normalize_text("\n".join(parts))
    if not text:
        raise ValueError("No extractable text found in the DOCX file.")
    metadata = {
        "paragraphs": sum(1 for paragraph in document.paragraphs if paragraph.text.strip()),
        "tables": len(document.tables),
    }
    return ParsedDocument(text=text, extension=".docx", metadata=metadata)


def extract_doc_text(payload: bytes) -> ParsedDocument:
    """Extract text from a legacy .doc file using binary pattern matching."""
    try:
        from docx import Document
        doc = Document(BytesIO(payload))
        parts: List[str] = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        text = normalize_text("\n".join(parts))
        if text:
            return ParsedDocument(text=text, extension=".doc")
    except Exception:
        pass

    try:
        import re as _re
        text_chunks: List[str] = []
        for encoding in ("utf-16-le", "utf-16-be", "latin-1", "cp1252"):
            try:
                decoded = payload.decode(encoding, errors="replace")
                words = _re.findall(r'[\w\s.,;:!?()\-\'"]+', decoded)
                readable = " ".join(w.strip() for w in words if len(w.strip()) > 2)
                if len(readable) > 100:
                    text_chunks.append(readable)
            except Exception:
                continue
        if text_chunks:
            best = max(text_chunks, key=len)
            text = normalize_text(best)
            if text:
                return ParsedDocument(
                    text=text,
                    extension=".doc",
                    warnings=["Legacy .doc file: extracted via binary fallback. Some formatting may be lost."],
                )
    except Exception:
        pass

    raise ValueError(
        "Could not extract text from this .doc file. "
        "Please save it as .docx or PDF and try again."
    )


def extract_image_text(payload: bytes, extension: str, api_manager: Optional[APIManager]) -> ParsedDocument:
    """Extract resume text from an image through Kimi Vision when configured."""
    if api_manager is None:
        raise ValueError("Image parsing requires a configured Kimi API manager.")
    image_url = image_payload_to_data_url(payload, extension)
    provider = api_manager.preferred_chat_provider()
    content = [
        {"type": "image_url", "image_url": {"url": image_url}},
        {
            "type": "text",
            "text": "Extract all visible resume text. Return plain text only.",
        },
    ]
    response_text = api_manager.chat_completion(
        messages=[
            {"role": "system", "content": "You extract resume text accurately and never add content that is not visible."},
            {"role": "user", "content": content},
        ],
        provider=provider,
        model=api_manager.provider_model(provider),
        max_tokens=1800,
        temperature=0.0,
        extra_body={"thinking": {"type": "disabled"}},
    )
    text = normalize_text(response_text)
    if not text:
        raise ValueError("Kimi Vision returned no extractable text.")
    return ParsedDocument(text=text, extension=extension, warnings=["Image text was extracted through Kimi Vision."])


def image_payload_to_data_url(payload: bytes, extension: str) -> str:
    """Convert image bytes to a base64 data URL accepted by vision chat APIs."""
    mime = "jpeg" if extension.lower() in {".jpg", ".jpeg"} else extension.lower().lstrip(".")
    encoded = base64.b64encode(payload).decode("utf-8")
    return f"data:image/{mime};base64,{encoded}"


def normalize_text(value: str) -> str:
    """Normalize whitespace while preserving useful line breaks."""
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in value.splitlines()]
    cleaned = "\n".join(line for line in lines if line)
    return re.sub(r"\n{3,}", "\n\n", cleaned).strip()


def has_multicolumn_layout(words: Iterable[Dict[str, float]]) -> bool:
    """Detect likely multi-column layout from word x-coordinate clusters."""
    x_positions = sorted(float(word.get("x0", 0.0)) for word in words)
    if len(x_positions) < 30:
        return False
    midpoint = len(x_positions) // 2
    left_cluster = x_positions[:midpoint]
    right_cluster = x_positions[midpoint:]
    if not left_cluster or not right_cluster:
        return False
    gap = min(right_cluster) - max(left_cluster)
    return gap > 80


def detect_hidden_text(text: str) -> str:
    """Flag suspicious text streams that often indicate ATS-hostile hidden content."""
    repeated_space_ratio = text.count("  ") / max(len(text), 1)
    if repeated_space_ratio > 0.05:
        return "Suspicious spacing detected; verify there is no hidden or layered text."
    return ""


def dedupe(values: Iterable[str]) -> List[str]:
    """Return values in original order without duplicates."""
    seen: set[str] = set()
    output: List[str] = []
    for value in values:
        if value and value not in seen:
            seen.add(value)
            output.append(value)
    return output
